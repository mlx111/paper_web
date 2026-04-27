"""Lightweight on-demand document text extraction service."""

from __future__ import annotations

import re
from html.parser import HTMLParser
from pathlib import Path
from typing import Any


def _ok(**payload: Any) -> dict[str, Any]:
    return {"ok": True, **payload}


def _err(code: str, message: str) -> dict[str, Any]:
    return {"ok": False, "error": code, "message": message}


class _TextHTMLParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self._skip_depth = 0
        self.parts: list[str] = []

    def handle_starttag(self, tag: str, attrs):
        if tag.lower() in {"script", "style", "noscript"}:
            self._skip_depth += 1
        if tag.lower() in {"p", "br", "div", "section", "article", "h1", "h2", "h3", "li"}:
            self.parts.append("\n")

    def handle_endtag(self, tag: str):
        if tag.lower() in {"script", "style", "noscript"} and self._skip_depth:
            self._skip_depth -= 1
        if tag.lower() in {"p", "div", "section", "article", "li"}:
            self.parts.append("\n")

    def handle_data(self, data: str):
        if not self._skip_depth and data.strip():
            self.parts.append(data)

    def text(self) -> str:
        return _normalize_text(" ".join(self.parts))


def _normalize_text(text: str) -> str:
    text = text.replace("\x00", "")
    text = re.sub(r"[ \t\r\f\v]+", " ", text)
    text = re.sub(r"\n\s*\n\s*", "\n\n", text)
    return text.strip()


class DocumentParserService:
    """Extract text for ad-hoc Agent tool use.

    This service is intentionally separate from the upload -> split -> index
    pipeline. It is meant for quick file reads during a conversation.
    """

    supported_extensions = {".txt", ".text", ".md", ".markdown", ".html", ".htm", ".pdf", ".docx"}

    def __init__(self, allowed_roots: list[Path] | None = None):
        project_root = Path(__file__).resolve().parents[2]
        self.allowed_roots = [
            root.resolve()
            for root in (
                allowed_roots
                if allowed_roots is not None
                else [
                    project_root / "uploads",
                    project_root / "app" / "data",
                    project_root / "workspace",
                    project_root,
                ]
            )
        ]

    def extract_text_from_file(self, file_path: str, summary_length: int = 5000) -> dict[str, Any]:
        if not file_path or not isinstance(file_path, str):
            return _err("INVALID_INPUT", "file_path must be a non-empty string")

        path = Path(file_path).expanduser()
        if not path.is_absolute():
            path = (Path.cwd() / path).resolve()
        else:
            path = path.resolve()

        if not self._is_allowed(path):
            return _err("PATH_NOT_ALLOWED", f"path is outside allowed roots: {path}")
        if not path.exists() or not path.is_file():
            return _err("FILE_NOT_FOUND", f"file not found: {path}")

        suffix = path.suffix.lower()
        if suffix not in self.supported_extensions:
            return _err("UNSUPPORTED_FILE_TYPE", f"unsupported file type: {suffix}")

        try:
            if suffix in {".txt", ".text", ".md", ".markdown"}:
                text = self._read_text(path)
            elif suffix in {".html", ".htm"}:
                text = self._read_html(path)
            elif suffix == ".pdf":
                text = self._read_pdf(path)
            elif suffix == ".docx":
                text = self._read_docx(path)
            else:
                return _err("UNSUPPORTED_FILE_TYPE", f"unsupported file type: {suffix}")
        except Exception as exc:
            return _err("PARSE_FAILED", f"failed to parse {path.name}: {exc}")

        summary = self._summarize(text, summary_length)
        return _ok(
            file_path=str(path),
            filename=path.name,
            file_type=suffix,
            text=text,
            summary=summary,
            char_count=len(text),
        )

    def _is_allowed(self, path: Path) -> bool:
        for root in self.allowed_roots:
            try:
                path.relative_to(root)
                return True
            except ValueError:
                continue
        return False

    @staticmethod
    def _read_text(path: Path) -> str:
        return _normalize_text(path.read_text(encoding="utf-8", errors="ignore"))

    @staticmethod
    def _read_html(path: Path) -> str:
        parser = _TextHTMLParser()
        parser.feed(path.read_text(encoding="utf-8", errors="ignore"))
        return parser.text()

    @staticmethod
    def _read_pdf(path: Path) -> str:
        try:
            import pdfplumber
        except ImportError:
            pdfplumber = None

        if pdfplumber is not None:
            parts: list[str] = []
            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        parts.append(page_text)
            return _normalize_text("\n\n".join(parts))

        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("pdfplumber or pypdf is required for PDF parsing") from exc

        reader = PdfReader(str(path))
        return _normalize_text("\n\n".join(page.extract_text() or "" for page in reader.pages))

    @staticmethod
    def _read_docx(path: Path) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx is required for DOCX parsing") from exc

        doc = Document(str(path))
        parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return _normalize_text("\n".join(parts))

    @staticmethod
    def _summarize(text: str, summary_length: int) -> str:
        summary_length = max(1, min(int(summary_length or 5000), 20000))
        if len(text) <= summary_length:
            return text
        return text[:summary_length].rstrip() + "..."


document_parser_service = DocumentParserService()
